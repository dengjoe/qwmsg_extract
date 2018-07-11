#!/usr/bin/env python3
# -*- coding: utf-8 -*-


"""
提取qq的消息文本中: 
	关注的人的所有发言，如果是发言内容有：@xxx 的，则反向将xxx的话也提取出来，形成问答格式。
"""

__author__ = 'kevin deng'

import os
import sys
import re
import getopt
import time
import datetime

import qqmsg_db
import qqmsg_item
import qmextract_ini
import global_data
import pdb

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# 配置文件的部分内容
important_names = ["刘明慧", "山长 清一"]
#section_wx = "wx_config"

filename_db    = "msgs.db"
error_filename = "error_msg.log"
errname_filename = "error_name.log"



# 判断是否关注人的消息，如果是，则搜索内容是否为带@的应答，如果是应答，则找回原先的问题。
# 将问题和应答存入输出文件。如果找不到应答，信息全部存入错误日志
last_time = ""

def output_important_msg(db, fout, outype, msg):
	""" 提取重要的人的应答消息，并输出到文件 """
	global last_time
	if msg.time[0:10] > last_time:
		print("last:",last_time, "now:", msg.time[0:10])
		output_msg_date(fout, outype, msg)

	last_time = msg.time[0:10]

	if msg.name in important_names:
		m = re.match(r'(@\w?\d{0,4}[\w\+\-\s]?[\w\d\_\-\—\*\^\.\,\~﹏！\=\{\}\(\)、\!\"\'\[\]]{1,15}[\+\-\s]{0,3}\w{0,6})\s(.+)', msg.content)
		if m:
			# 是含有应答对象的内容，提取出对象的姓名
			name = m.group(1).strip()[1:]
			# 山长的消息另外单独处理
			if msg.name == "山长 清一":
				keyname = get_keyname(db, name)
				shortname = re.sub(r"\d{0,4}", "", keyname).strip()
				msg.content = re.sub("@"+name, "@"+shortname, msg.content)
				# print(name, keyname)
				name = keyname

			# 只记录2小时（7200秒）内的回复信息
			msg_last = db.get_last_msg_by_name(name)
			if msg_last and msg_time_diff(msg, msg_last) < 7200:
				output_msg(fout, outype, msg_last)
			else:
				log_errmsg_at_name(msg, name)

		output_msg(fout, outype, msg, "\n")
		global_data.output_msg_count += 1

def get_keyname(db, name):
	# pdb.set_trace()
	names = db.get_keynames(name)
	if names:
		if len(names)>1:
			print(names) # only for help
		return names[0]
	else:
		return name
		
def msg_time_diff(msg, msg_last):
	tmlast = datetime.datetime.strptime(msg_last.time, "%Y-%m-%d %H:%M:%S")
	tmnow = datetime.datetime.strptime(msg.time, "%Y-%m-%d %H:%M:%S")
	return (tmnow-tmlast).seconds

def docx_add_style_run(p, text, fontsize, isbold, fontname=None):
		run = p.add_run(text)
		run.font.size = Pt(fontsize)
		run.font.bold = isbold
		if fontname:
			run.font.name = fontname
			run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
		return run

def output_msg_date(doc, outype, msg):
	if outype == 0:
		print(msg.time[0:10] + "\n\n")
	elif outype == 1:
		doc.write(msg.time[0:10] + "\n\n")
	elif outype == 2:
		p = doc.add_heading('', level=1)
		docx_add_style_run(p, msg.time[0:10], 12, True, "宋体")
		doc.add_paragraph('')


def output_msg(doc, outype, msg, append=None):
	iname = re.sub(r"\d{0,4}", "", msg.name).strip()
	if iname.find("刘明慧") >= 0:
		iname = "刘老师"

	strout = iname + "：" + msg.content + "\n"
	if append:
		strout = strout + append

	try:
		if outype == 0:
			print(strout)
		elif outype == 1:
			doc.write(strout)
		elif outype == 2:
			p = doc.add_paragraph('')
			run = docx_add_style_run(p, iname+ "：", 12, True, "宋体")
			if append:
				run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
			docx_add_style_run(p, msg.content, 12, False, "宋体")

			if append:
				p =doc.add_paragraph('')
				docx_add_style_run(p, "", 12, False, "宋体")

	except:
		print("err content:", "["+ msg.time +"]" + msg.name + "：" + msg.content)	



def log_errmsg_at_name(msg, name):
	""" 记录姓名找不到的错误日志文件 """
	global_data.error_msg_count += 1
	with open(error_filename, 'a') as f:
		f.write(msg.name + " [" + msg.time + "] " + "error name: " + name + "\n")
		f.write(msg.content + "\n")


def log_run_time(filename):
	""" 记录运行时间到日志文件 """
	with open(filename, 'a', encoding='utf8') as f:
		prefix = " --------------- "
		f.write("\n" + prefix + global_data.runtime.strftime('%Y-%m-%d %H:%M:%S') + prefix +"\n\n")


#解析QQ消息-->DB
def parse_msg(db, fout, strlines, outype):
	""" 从字符串列表中提取出逐条信息，格式化后，存入数据库 """

	log_run_time(error_filename)

	#读取最后一次写入的消息时间
	skip_flag = 0
	last_msg = db.get_last_msg()
	if last_msg:
		skip_flag = 1
		old_time = last_msg.timestamp
		print("oldtime:", last_msg.time)

	#逐行提取消息
	pat =       re.compile(r'^\d{4}\-\d{2}\-\d{2}\s\d{1,2}:\d{1,2}:\d{1,2}\s')
	pat_group = re.compile(r'(^\d{4}\-\d{2}\-\d{2}\s\d{1,2}:\d{1,2}:\d{1,2}) (.+)')
	msg = qqmsg_item.MsgItem(errname_filename)

	for line in strlines:
		if pat.match(line):
			m = pat_group.match(line)
			# print(line)
			# print("1:",m.group(1), "2:", m.group(2))

			# 对比时间，跳过已经处理过的消息
			time_array = time.strptime(m.group(1), "%Y-%m-%d %H:%M:%S")
			now_time = int(time.mktime(time_array))

			if skip_flag==1 and old_time>now_time:
				continue

			if msg.time !="":
				output_important_msg(db, fout, outype, msg)
				db.save_msg(msg)

			msg = qqmsg_item.MsgItem(errname_filename)
			msg.set_time(m.group(1))
			msg.set_name(m.group(2))
		else:
			str1 = line.strip()
			if len(str1)>1:
				msg.add_content(str1)

	#记录最后一条消息
	output_important_msg(db, fout, outype, msg)
	db.save_msg(msg)



#0、打开数据库，获取最后一条信息的时间和内容，对比QQ记录，找到对应时间的下一条记录，从这里开始
#1、先按行开头的日期时间，来确定一条记录。并提取本行后面的昵称\时间，记录到自己的类，保存到数据库。
#2、如果含有关注名的，则一定要记录并执行3、4
#3、搜索后续内容中是否有'\@\w+\s'，有则提取出该名字
#4、在之前记录中搜索该名字的最近一条记录，并写入在本条之前。如果没有找到，报一错误，并记录当前时间和内容
#5、如果3、4有结果，先写入4的记录，再记录本条记录。否则只记录本条记录

def file_extension(path):
	return os.path.splitext(path)[1]  

def qqmsg_extract(dbname, outputname, inputname):
	"""提取群消息"""
	db = qqmsg_db.Qqmsg_db(dbname, errname_filename)
	fin = open(inputname, 'r', encoding='utf8', errors='ignore')

	outype = 0 # 0-stdout, 1-txt, 2-docx
	outitle = "群记录"
	file_ext = file_extension(outputname)
	if file_ext == ".txt":
		fout = open(outputname, 'w')
		fout.write(outitle + "\n\n")
		outype = 1
	elif file_ext == ".docx":
		fout = Document()
		p = fout.add_paragraph("")
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # CENTER居中对齐 LEFT左对齐 RIGHT右对齐
		run = docx_add_style_run(p, outitle, 18, True, "宋体")
		fout.add_paragraph("")	
		outype = 2
	else:
		print("stdout msg!")

	try:
		parse_msg(db, fout, fin.readlines(), outype)
	finally:
		if fin:
			fin.close()
		if file_ext == ".txt" and fout:
			fout.close()
		if file_ext == ".docx" and fout:
			fout.save(outputname)

def qqmsg_save_nicknames(dbname, nickfile):
	db = qqmsg_db.Qqmsg_db(dbname, errname_filename)
	pat = re.compile(r'([.\d\w\s\-\—\_\+\，·]*)\<([.\d\w\s\-\—\_\+\*\^\.\,\~﹏！\!\=\{\}\(\)、\"\'\[\]（）]*)')
	with open(nickfile, 'r', encoding='utf8') as f:
		for line in f.readlines():
			m = pat.match(line)
			if m:
				#nickname = m.group(2)[:len(m.group(2))-1]
				db.add_nickname(m.group(1).strip(), m.group(2).strip())
			else:
				print("error in nickname:", line)
	#db.stdout_nickname()

def qqmsg_clear_nicknames(dbname):
	db = qqmsg_db.Qqmsg_db(dbname, errname_filename)
	db.clear_nickname()


def qqmsg_init():
	""" 定位当前目录，并重新确定各配置和出错的文件路径，并读取配置文件 """
	global filename_db
	global error_filename
	global errname_filename
	path = os.getcwd() #sys.path[0]
	filename_db      = path + "/" + filename_db
	error_filename   = path + "/" + error_filename
	errname_filename = path + "/" + errname_filename

	# 读取配置文件
	global important_names
	filename_ini  = path + "/" + "qmextract.ini"
	section       = "qq_config"
	print(path, filename_ini)

	meini = qmextract_ini.Qqmsg_ini(filename_ini, section)
	important_names = meini.get_names()
	return meini

def extract_result():
	str_result = "\n ---------------------------------- \n" \
	+ "开始时间：   " + global_data.runtime.strftime('%Y-%m-%d %H:%M:%S') + "\n" \
	+ "无法识别的名字数量：    " + str(global_data.error_name_count) + "\n" \
	+ "无法找到提问的消息数量：" + str(global_data.error_msg_count) + "\n" \
	+ "提取消息数量：          " + str(global_data.output_msg_count) + "\n" 
	return str_result


def main():
	args = sys.argv
	usage = args[0] + " -i inputname -o outputname"
	if len(args)<5:
		print(usage)
		return -1

	opts, args = getopt.getopt(sys.argv[1:], "i:o:")
	input_file=""
	output_file=""

	for op, value in opts:
		if op == "-i":
			input_file = value
		elif op == "-o":
			output_file = value

	if input_file == "" or output_file == "":
		print(usage)
		return -2


	# 提取qq消息，保存并导出关注的内容
	qqmsg_init()
	global_data.init_global_data()

	qqmsg_extract(filename_db, output_file, input_file)
	#qqmsg_extract(filename_db, "./out2.txt","./清粉家园群02.txt")

	# 输出运行统计
	print(extract_result())




if __name__ == '__main__':
	main()	